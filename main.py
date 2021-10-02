import io
from tkinter import *
from tkinter import filedialog
import re

import docx
import numpy as np
from tkinter import ttk
from datetime import datetime
import tkinter as tk
from docx.shared import Inches
from mysql.connector import Binary
from resizeimage import resizeimage
from PIL import Image, ImageTk
from keras.models import Sequential, model_from_json
from keras.layers import Conv2D, MaxPooling2D, Flatten
from keras.layers import Dense, Dropout
from keras.preprocessing.image import ImageDataGenerator
import pyodbc
from keras_preprocessing import image

connection = pyodbc.connect("Driver={SQL Server Native Client 11.0};"
                               "Server=ALEXANDR\SQLEXPRESS01;"
                               "Database=Med;"
                               "Trusted_Connection=yes;")

class Main(tk.Frame):
    def __init__(self, root):
        super().__init__(root)
        self.init_main()
        self.view_records()

    def init_main(self):

        toolbar = Frame(height=700, width=1300)
        toolbar.place(x=0,y=0)
        newbar = Frame(height= 100,width=1300)
        newbar.place(x=0,y=0)
        newbar2 = Frame(height=600, width=1300)
        newbar2.place(x=0, y=100)
        btn_open_dialog = tk.Button(newbar, text='Провести диагностику', command=self.open_form, font=("Verdana", 10),
                               background='#DCDCDC',width =27)
        btn_open_dialog.place(x=5,y=60)

        if(user.admin==1):
            btn_open_dialog = tk.Button(newbar, text='Управление пользователями',
                                        command=self.manage,font=("Verdana", 10),
                                        background='#DCDCDC',width =27)
            btn_open_dialog.place(x=235, y=60)

        btn_open_dialog = tk.Button(newbar, text='Применить фильтры',
                                    command=self.filter, font=("Verdana", 10),
                                    background='#DCDCDC',width =27)
        btn_open_dialog.place(x=1065, y=8)

        btn_open_dialog = tk.Button(newbar, text='Сбросить фильтры',
                                    command=self.reset, font=("Verdana", 10),
                                    background='#DCDCDC',width =27)
        btn_open_dialog.place(x=1065, y=60)

        btn_open_dialog = tk.Button(newbar, text='Данные выбранного пациента',
                                    command=self.find_pat, font=("Verdana", 10),
                                    background='#DCDCDC', width =27)
        btn_open_dialog.place(x=630, y=60)

        self.filt1 = ttk.Combobox(newbar,
                                             values=['Все состояния','COVID','NORMAL', 'PNEUMONIA',
                                                     'TUBERCULOSIS'], width=29)
        self.filt1.place(x=850, y=12)
        self.filt1.current(0)

        self.filt2 = ttk.Combobox(newbar,
                                             values=['За всё время','За один день','За неделю',
                                                     'За месяц', 'За полгода'], width=29)
        self.filt2.place(x=650, y=12)
        self.filt2.current(0)

        self.filt3 = ttk.Combobox(newbar,
                                             values=['Все результаты','Совпадение диагнозов',
                                                     'Несовпадение диагнозов', 'Диагноз не поставлен'], width=29)
        self.filt3.place(x=450, y=12)
        self.filt3.current(0)

        label_main = tk.Label(newbar, text='Пользователь:', font=("Verdana", 15, 'bold'))
        label_main.place(x=5, y=5)

        label_main = tk.Label(newbar, text=user.Surname+" "+user.Name, font=("Verdana", 13, 'bold'))
        label_main.place(x=180, y=8)

        self.entry_i = Entry(newbar)
        self.entry_i.configure(font=("Verdana", 10), width=15)
        self.entry_i.place(x=500, y=63)

        self.tree = ttk.Treeview(newbar2, columns=( 'ID','FIO', 'Date_of_birth',
                                                    'Polis','Own1', 'Sys_res','Usr_res', 'Res', 'Date'),
                                 height=29, show='headings')
        self.tree.column("ID", width=30, anchor=tk.W)
        self.tree.column("FIO", width=200, anchor=tk.W)
        self.tree.column("Date_of_birth", width=100, anchor=tk.W)
        self.tree.column("Polis", width=100, anchor=tk.W)
        self.tree.column("Own1", width=160, anchor=tk.W)
        self.tree.column("Sys_res", width=200, anchor=tk.W)
        self.tree.column("Usr_res", width=180, anchor=tk.W)
        self.tree.column("Res", width=180, anchor=tk.W)
        self.tree.column("Date", width=125, anchor=tk.W)

        self.tree.heading("ID", text='№')
        self.tree.heading("FIO", text='ФИО пациента')
        self.tree.heading("Date_of_birth", text='Дата рождения')
        self.tree.heading("Polis", text='Полис')
        self.tree.heading("Own1", text='Провёл диагностику')
        self.tree.heading("Sys_res", text='Диагноз системы')
        self.tree.heading("Usr_res", text='Поставил диагноз')
        self.tree.heading("Res", text='Диагноз врача')
        self.tree.heading("Date", text='Дата проведения')

        self.tree.bind('<<TreeviewSelect>>', self.click)
        self.tree.bind('<Double-Button-1>', self.add_res)
        self.tree.pack(side=LEFT)

        scroll = tk.Scrollbar(newbar2, command=self.tree.yview)
        scroll.pack(side=LEFT,fill=tk.Y)
        self.tree.configure(yscrollcommand=scroll.set)

    def open_form(self):
        Child()

    def reset(self):
        self.view_records()
        self.filt1.current(0)
        self.filt2.current(0)
        self.filt3.current(0)

    def click(self,event):

        self.entry_i.delete(0, 'end')
        curItem = self.tree.focus()
        self.l =self.tree.item(curItem)
        self.entry_i.insert(0,self.l['values'][3])

    def find_pat(self):
        dbCursor = connection.cursor()
        requestString = f'''select id,[name],surname,patronymic,date_of_birth,polis,sys_res,res_r,[date],
                usr_name,usr_sorname,usr_res
                 from [Patients] join Test on Patients.pat_id = Test.pat_id
        		 join Usr on Usr.usr_id = Test.usr_id 
        		 where polis = {self.l['values'][3]}'''
        dbCursor.execute(requestString)
        [self.tree.delete(i) for i in self.tree.get_children()]
        for row in dbCursor:
            self.tree.insert('', 'end', values=(
            row[0], row[2] + " " + row[1] + " " + row[3], row[4], row[5], row[10] + " " + row[9],
            row[6], row[11], row[7], row[8]))

    def manage(self):
        Manage_usr()

    def view_records(self):
        dbCursor = connection.cursor()
        requestString='''select id,[name],surname,patronymic,date_of_birth,polis,sys_res,res_r,[date],
        usr_name,usr_sorname,usr_res
         from [Patients] join Test on Patients.pat_id = Test.pat_id
		 join Usr on Usr.usr_id = Test.usr_id '''
        dbCursor.execute(requestString)
        [self.tree.delete(i) for i in self.tree.get_children()]
        for row in dbCursor:
            self.tree.insert('', 'end', values=(row[0],row[2]+" "+row[1]+" "+row[3],row[4],row[5],row[10]+" "+row[9],
                                                row[6],row[11],row[7],row[8]))

    def add_res(self,event):
        selected_item = self.tree.selection()[0]
        self.values = self.tree.item(selected_item, option="values")
        ch = Show_res()

    def filter(self):
        unit_to_multiplier1 = {
            'Все состояния': '%',
            'COVID': 'COVID',
            'NORMAL': 'NORMAL',
            'PNEUMONIA': 'PNEUMONIA',
            'TUBERCULOSIS': 'TUBERCULOSIS'
        }
        mult1 = unit_to_multiplier1[self.filt1.get()]

        unit_to_multiplier2 = {
            'За всё время': '%',
            'За один день': 'day,-1',
            'За неделю': 'Week,-1',
            'За месяц': 'month,-1',
            'За полгода': 'month,-6'
        }
        mult2 = unit_to_multiplier2[self.filt2.get()]

        if mult2 =='%':
            s = ''
        else:
            s = f"and date > DATEADD({mult2},GETDATE())"

        unit_to_multiplier3 = {
            'Все результаты': '',
            'Совпадение диагнозов': 'and res_r=sys_res',
            'Несовпадение диагнозов': 'and res_r <> sys_res',
            'Диагноз не поставлен': 'and res_r is null',
        }
        mult3 = unit_to_multiplier3[self.filt3.get()]

        dbCursor = connection.cursor()
        requestString = f'''select id,[name],surname,patronymic,date_of_birth,polis,sys_res,res_r,[date],
                usr_name,usr_sorname,usr_res
                 from [Patients] join Test on Patients.pat_id = Test.pat_id
        		 join Usr on Usr.usr_id = Test.usr_id
        		 where (sys_res LIKE '{mult1}' or res_r LIKE '{mult1}') {s} {mult3}'''
        dbCursor.execute(requestString)
        [self.tree.delete(i) for i in self.tree.get_children()]
        for row in dbCursor:
            self.tree.insert('', 'end', values=(
            row[0], row[2] + " " + row[1] + " " + row[3], row[4], row[5], row[10] + " " + row[9],
            row[6], row[11], row[7], row[8]))

class Manage_usr(tk.Toplevel):
    def __init__(self):
        super().__init__(root)
        self.init_child()
        self.view_records()

    def init_child(self):
        self.title('Управление пользователями')
        self.geometry('1140x720+10+50')
        self.resizable(False, False)

        self.toolbar = Frame(self,height=700, width=1300)
        self.toolbar.place(x=0, y=0)
        self.newbar = Frame(self,height=100, width=1300)
        self.newbar.place(x=0, y=0)
        self.newbar2 = Frame(self,height=600, width=1300)
        self.newbar2.place(x=0, y=100)

        btn_open_dialog1 = tk.Button(self.newbar, text='Добавить нового пользователя', command=self.new_usr, font=("Verdana", 10),
                                    background='#DCDCDC',width =30)
        btn_open_dialog1.place(x=5, y=60)

        btn_open_dialog2 = tk.Button(self.newbar, text='Удалить выбранного пользователя', command=self.destr,
                                    font=("Verdana", 10),
                                    background='#DCDCDC',width =30)
        btn_open_dialog2.place(x=325, y=60)

        self.tree = ttk.Treeview(self.newbar2, columns=('ID', 'FIO', 'Date_of_birth',
                                                   'Polis', 'Own1', 'Sys_res',),
                                 height=29, show='headings')
        self.tree.column("ID", width=30, anchor=tk.W)
        self.tree.column("FIO", width=120, anchor=tk.W)
        self.tree.column("Date_of_birth", width=120, anchor=tk.W)
        self.tree.column("Polis", width=120, anchor=tk.W)
        self.tree.column("Own1", width=120, anchor=tk.W)
        self.tree.column("Sys_res", width=60, anchor=tk.W)

        self.tree.heading("ID", text='№')
        self.tree.heading("FIO", text='Фамилия')
        self.tree.heading("Date_of_birth", text='Имя')
        self.tree.heading("Polis", text='Логин')
        self.tree.heading("Own1", text='Пароль')
        self.tree.heading("Sys_res", text='Роль')

        self.tree.pack(side=LEFT)

        scroll = tk.Scrollbar(self.newbar2, command=self.tree.yview)
        scroll.pack(side=LEFT,fill=tk.Y)
        self.tree.configure(yscrollcommand=scroll.set)

        self.tree1 = ttk.Treeview(self.newbar2, columns=('ID', 'FIO', 'Date_of_birth',
                                                        'Polis'),
                                 height=29, show='headings')
        self.tree1.column("ID", width=30, anchor=tk.W)
        self.tree1.column("FIO", width=200, anchor=tk.W)
        self.tree1.column("Date_of_birth", width=180, anchor=tk.W)
        self.tree1.column("Polis", width=120, anchor=tk.W)

        self.tree1.heading("ID", text='№')
        self.tree1.heading("FIO", text='Пользователь')
        self.tree1.heading("Date_of_birth", text='Действие')
        self.tree1.heading("Polis", text='Дата и время')

        self.tree1.pack(side=LEFT)

        scroll = tk.Scrollbar(self.newbar2, command=self.tree1.yview)
        scroll.pack(side=LEFT, fill=tk.Y)
        self.tree1.configure(yscrollcommand=scroll.set)

        self.grab_set()

    def view_records(self):
        dbCursor = connection.cursor()
        requestString='''select * from Usr'''
        dbCursor.execute(requestString)
        [self.tree.delete(i) for i in self.tree.get_children()]
        for row in dbCursor:
            if row[5] == 1:
                rol = 'Admin'
            else:
                rol= 'User'
            self.tree.insert('', 'end', values=(row[0],row[2],row[1],row[3],row[4],rol))

        requestString = '''select log_id,action,act_date,usr_name,usr_sorname
         from Logr join Usr on Usr.usr_id = Logr.usr_id'''
        dbCursor.execute(requestString)
        [self.tree1.delete(i) for i in self.tree1.get_children()]
        for row in dbCursor:
            self.tree1.insert('', 'end', values=(row[0], row[4]+" "+row[3], row[1], row[2]))
        self.focus_set()
        connection.commit()

    def new_usr(self):
        New_usr(self.tree,self.tree1)

    def destr(self):
        dbCursor = connection.cursor()
        selected_item = self.tree.selection()[0]
        self.values = self.tree.item(selected_item, option="values")
        requestString = '''delete from Usr where usr_id =? '''
        dbCursor.execute(requestString,(self.values[0]))
        connection.commit()
        bd.logr("Удалил пользователя")
        self.view_records()

class New_usr(tk.Toplevel):
    def __init__(self,tree,tree1):
        super().__init__(root)
        self.init_main()
        self.tree =tree
        self.tree1 =tree1

    def init_main(self):
        self.title("Добавление пользователя")
        self.geometry("400x300+10+50")
        self.resizable(False, False)

        label_main = tk.Label(self, text='Данные пользователя:', font=("Verdana", 15, 'bold'))
        label_main.place(x=50, y=10)

        label_description = tk.Label(self, text='Фамилия:', font=("Verdana", 10))
        label_description.place(x=30, y=80)
        label_select = tk.Label(self, text='Имя:', font=("Verdana", 10))
        label_select.place(x=30, y=110)
        label_sum = tk.Label(self, text='Логин:', font=("Verdana", 10))
        label_sum.place(x=30, y=140)
        label_sum2 = tk.Label(self, text='Пароль:', font=("Verdana", 10))
        label_sum2.place(x=30, y=170)
        label_sum2 = tk.Label(self, text='Роль:', font=("Verdana", 10))
        label_sum2.place(x=30, y=200)

        self.entry_name = Entry(self)
        self.entry_name.configure(font=("Verdana", 10),width=24)
        self.entry_name.place(x=120, y=80)

        self.entry_sorname = Entry(self)
        self.entry_sorname.configure(font=("Verdana", 10),width=24)
        self.entry_sorname.place(x=120, y=110)

        self.entry_patronymic = Entry(self)
        self.entry_patronymic.configure(font=("Verdana", 10),width=24)
        self.entry_patronymic.place(x=120, y=140)

        self.entry_data = Entry(self)
        self.entry_data.configure(font=("Verdana", 10),width=24)
        self.entry_data.place(x=120, y=170)

        self.selfcomboExample = ttk.Combobox(self,
                                    values=["Admin","User",],width=29)
        self.selfcomboExample.place(x=120,y=200)
        self.selfcomboExample.current(1)

        btn_cancel = tk.Button(self, text='Добавить', command=self.insert_usr, font=("Verdana", 10),
                               background='#DCDCDC')
        btn_cancel.place(x=30, y=250)

        btn_cancel = tk.Button(self, text='Закрыть окно', command=self.destroy, font=("Verdana", 10),
                               background='#DCDCDC')
        btn_cancel.place(x=215, y=250)

        self.grab_set()
        self.focus_set()

    def insert_usr(self):
        if self.selfcomboExample.get() == 'Admin':
            self.l=1
        else:
            self.l=0
        dbCursor = connection.cursor()
        requestString = "insert into Usr(usr_name,usr_sorname,usr_login,usr_password,admin) " \
                        "values(?,?,?,?,?);"
        dbCursor.execute(requestString, (self.entry_sorname.get(),
                                         self.entry_name.get(),
                                         self.entry_patronymic.get(),
                                         self.entry_data.get(),
                                         self.l))
        connection.commit()
        bd.logr("Добавил пользователя")
        self.view_records()

    def view_records(self):
        dbCursor = connection.cursor()
        requestString='''select * from Usr'''
        dbCursor.execute(requestString)
        [self.tree.delete(i) for i in self.tree.get_children()]
        for row in dbCursor:
            if row[5] == 1:
                rol = 'Admin'
            else:
                rol= 'User'
            self.tree.insert('', 'end', values=(row[0],row[2],row[1],row[3],row[4],rol))

        requestString = '''select log_id,action,act_date,usr_name,usr_sorname
                 from Logr join Usr on Usr.usr_id = Logr.usr_id'''
        dbCursor.execute(requestString)
        [self.tree1.delete(i) for i in self.tree1.get_children()]
        for row in dbCursor:
            self.tree1.insert('', 'end', values=(row[0], row[4] + " " + row[3], row[1], row[2]))
        connection.commit()

class Show_res(tk.Toplevel):
    def __init__(self):
        super().__init__(root)
        self.view = app
        self.value = self.view.values[0]
        self.init_child()
        self.classes = ['COVID', 'NORMAL', 'PNEUMONIA', 'TUBERCULOSIS']

    def init_child(self):
        self.title('Тестирование')
        self.geometry('1100x700+10+50')
        self.resizable(False, False)

        label_main = tk.Label(self, text='Данные пациента:', font=("Verdana", 20, 'bold'))
        label_main.place(x=50, y=10)

        label_description = tk.Label(self, text='Фамилия:', font=("Verdana", 10))
        label_description.place(x=50, y=80)
        label_select = tk.Label(self, text='Имя:', font=("Verdana", 10))
        label_select.place(x=50, y=110)
        label_sum = tk.Label(self, text='Отчество:', font=("Verdana", 10))
        label_sum.place(x=50, y=140)
        label_sum2 = tk.Label(self, text='Дата рождения:', font=("Verdana", 10))
        label_sum2.place(x=50, y=170)
        label_sum3 = tk.Label(self, text='Полис:', font=("Verdana", 10))
        label_sum3.place(x=50, y=200)
        label_sum4 = tk.Label(self, text='Результат:', font=("Verdana", 10))
        label_sum4.place(x=50, y=340)
        label_sum5 = tk.Label(self, text='Диагноз врача:', font=("Verdana", 10))
        label_sum5.place(x=50, y=450)

        self.entry_name = Entry(self)
        self.entry_name.configure(font=("Verdana", 10))
        self.entry_name.place(x=200, y=80)

        self.entry_sorname = Entry(self)
        self.entry_sorname.configure(font=("Verdana", 10))
        self.entry_sorname.place(x=200, y=110)

        self.entry_patronymic = Entry(self)
        self.entry_patronymic.configure(font=("Verdana", 10))
        self.entry_patronymic.place(x=200, y=140)

        self.entry_data = Entry(self)
        self.entry_data.configure(font=("Verdana", 10))
        self.entry_data.place(x=200, y=170)

        self.entry_polis = Entry(self)
        self.entry_polis.configure(font=("Verdana", 10))
        self.entry_polis.place(x=200, y=200)

        self.entry_sys_res = Text(self, font=("Verdana", 10), width=39, height=5, wrap=WORD)
        self.entry_sys_res.place(x=50, y=360)

        self.selfcomboExample = ttk.Combobox(self,
                                             values=['COVID', 'NORMAL', 'PNEUMONIA', 'TUBERCULOSIS'], width=29)
        self.selfcomboExample.place(x=50, y=470)

        self.entry_res = Text(self, font=("Verdana", 10), width=39, height=8, wrap=WORD)
        self.entry_res.place(x=50, y=500)

        btn_cancel = tk.Button(self, text='Закрыть окно', command=self.destroy, font=("Verdana", 10),
                               background='#DCDCDC')
        btn_cancel.place(x=262, y=640)

        btn_save = tk.Button(self, text='Сохранить результат в файл', command=self.save_file, font=("Verdana", 10),
                               background='#DCDCDC')
        btn_save.place(x=155, y=240)

        btn_cancel1 = tk.Button(self, text='Сохранить диагноз', command=self.save_res, font=("Verdana", 10),
                               background='#DCDCDC')
        btn_cancel1.place(x=50, y=640)

        self.ImageFrame = LabelFrame(self, height=555, width=620)
        self.ImageFrame.place(x=400,y=80)

        self.show_res()

    def save_res(self):
        dbCursor = connection.cursor()
        requestString = "UPDATE Test SET res = ?,usr_res=?,res_r=? WHERE id= ? ;"
        dbCursor.execute(requestString, (self.entry_res.get("1.0",END),user.Surname+" "+user.Name,
                                         self.selfcomboExample.get(),self.value))
        connection.commit()
        app.view_records()
        bd.logr("Добавил диагноз")

    def show_res(self):
        unit_to_multiplier = {
            'COVID': 0,
            'NORMAL': 1,
            'PNEUMONIA': 2,
            'TUBERCULOSIS': 3,
        }
        dbCursor = connection.cursor()
        requestString = '''select [name],surname,patronymic,date_of_birth,polis,sys_res,res,[date],res_r,sys_confidence
         from [Patients] join Test on Patients.pat_id = Test.pat_id where id = ?;'''
        dbCursor.execute(requestString,(self.value))
        for row in dbCursor:
            self.entry_name.insert(0,row[1])
            self.entry_sorname.insert(0,row[0])
            self.entry_patronymic.insert(0,row[2])
            self.entry_data.insert(0,row[3])
            self.entry_polis.insert(0,row[4])
            self.entry_sys_res.insert(1.0,"Диагноз системы: "+row[5]+"\n\nУверенность системы: "+row[9])
            if row[6]!= None:
                self.entry_res.insert(1.0,row[6])
            if row[8]!= None:
                mult = unit_to_multiplier[row[8]]
                self.selfcomboExample.current(mult)
            self.test_id = self.value
            self.entry_name.configure(state = "disabled")
            self.entry_sorname.configure(state = "disabled")
            self.entry_patronymic.configure(state = "disabled")
            self.entry_data.configure(state = "disabled")
            self.entry_polis.configure(state = "disabled")
            self.entry_sys_res.configure(state = "disabled")

        requestString = '''SELECT * FROM IMG_Data where test_id=?'''
        dbCursor.execute(requestString,(self.test_id))

        for row in dbCursor:
            self.blob_img = row[1]
        img  =Image.open(io.BytesIO(self.blob_img))
        self.our_image2 = resizeimage.resize_contain(img, [555, 620])
        self.our_image2 = ImageTk.PhotoImage(self.our_image2)
        self.our_label2 = Label(self.ImageFrame, image=self.our_image2)
        self.our_label2.image = self.our_image2
        self.our_label2.pack()

    def save_file(self):
        self.Sfilename = filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[("Docx files", '*.docx')],
                title="Выбор имени файла")
        mydoc = docx.Document()
        mydoc.add_heading("Результат диагностики:", 0)

        dbCursor = connection.cursor()
        requestString = '''SELECT * FROM IMG_Data where test_id=?'''
        dbCursor.execute(requestString, (self.test_id))

        for row in dbCursor:
            self.blob_img = row[1]
        img = Image.open(io.BytesIO(self.blob_img))
        output = io.BytesIO()
        img.save(output, format='JPEG')
        output.seek(0)
        my_image = mydoc.add_picture(output,width=Inches(5.0))
        last_paragraph = mydoc.paragraphs[-1]
        last_paragraph.alignment = 1
        requestString = '''select [name],surname,patronymic,date_of_birth,polis,sys_res,res,[date],sys_confidence
                 from [Patients] join Test on Patients.pat_id = Test.pat_id where id = ?;'''
        dbCursor.execute(requestString, (self.value))
        for row in dbCursor:
            t = row[3]
            l = t.strftime('%d/%m/%Y')
            mydoc.add_heading("Пациент:", 2)
            mydoc.add_paragraph(row[1]+" "+row[0]+" "+row[2]+"."+"Дата рождения: "+l)
            mydoc.add_paragraph("Полис: "+self.entry_polis.get())
            mydoc.add_heading("Диагноз системы:", 2)
            mydoc.add_paragraph(row[5])
            mydoc.add_heading("Уверенность системы:", 2)
            mydoc.add_paragraph(row[8])
            mydoc.add_heading("Диагноз врача:", 2)
            mydoc.add_paragraph(row[6])
            mydoc.save(self.Sfilename)

class Child(tk.Toplevel):
    def __init__(self):
        super().__init__(root)
        self.init_child()
        self.view = app
        self.classes = ['COVID', 'NORMAL', 'PNEUMONIA', 'TUBERCULOSIS']

    def init_child(self):
        self.title('Тестирование')
        self.geometry('1100x700+10+50')
        self.resizable(False, False)


        label_main = tk.Label(self, text='Данные пациента:',font=("Verdana", 20, 'bold'))
        label_main.place(x=50, y=10)

        label_description = tk.Label(self, text='Фамилия:',font=("Verdana", 10))
        label_description.place(x=50, y=80)
        label_select = tk.Label(self, text='Имя:',font=("Verdana", 10))
        label_select.place(x=50, y=110)
        label_sum = tk.Label(self, text='Отчество:',font=("Verdana", 10))
        label_sum.place(x=50, y=140)
        label_sum2 = tk.Label(self, text='Дата рождения:', font=("Verdana", 10))
        label_sum2.place(x=50, y=170)
        label_sum3 = tk.Label(self, text='Полис:',font=("Verdana", 10))
        label_sum3.place(x=50, y=200)
        label_sum4 = tk.Label(self, text='Результат:', font=("Verdana", 10))
        label_sum4.place(x=50, y=340)
        label_sum5 = tk.Label(self, text='Диагноз врача:', font=("Verdana", 10))
        label_sum5.place(x=50, y=450)

        self.label_err1 = tk.Label(self,
                                  text='Такие поля как Имя, Фамилия,\nДата рождения и Полис\nне должны быть пустыми',
                                  font=("Verdana", 8), fg='#FF0000')

        self.label_err2 = tk.Label(self,
                                   text='Неверный формат даты!\nНужно День-Месяц-Год\nНапример(25-11-2000)',
                                   font=("Verdana", 8), fg='#FF0000')

        self.label_err3 = tk.Label(self,
                                   text='Поле полис должно содержать\nтолько цифры',
                                   font=("Verdana", 8), fg='#FF0000')

        self.entry_name = EntryWithPlaceholder(self,'Иванов')
        self.entry_name.configure(font=("Verdana", 10))
        self.entry_name.place(x=200, y=80)

        self.entry_sorname = EntryWithPlaceholder(self,'Иван')
        self.entry_sorname.configure(font=("Verdana", 10))
        self.entry_sorname.place(x=200, y=110)

        self.entry_patronymic = EntryWithPlaceholder(self,'Иванович')
        self.entry_patronymic.configure(font=("Verdana", 10))
        self.entry_patronymic.place(x=200, y=140)

        self.entry_data = EntryWithPlaceholder(self,'31-12-2000')
        self.entry_data.configure(font=("Verdana", 10))
        self.entry_data.place(x=200, y=170)

        self.entry_polis = EntryWithPlaceholder(self,'123456789012')
        self.entry_polis.configure(font=("Verdana", 10))
        self.entry_polis.place(x=200, y=200)

        self.entry_sys_res = Text(self, font=("Verdana", 10),width=39,height=5,wrap=WORD,state="disabled")
        self.entry_sys_res.place(x=50, y=360)

        self.selfcomboExample = ttk.Combobox(self,
                                             values=['COVID', 'NORMAL', 'PNEUMONIA', 'TUBERCULOSIS'], width=29)
        self.selfcomboExample.place(x=50, y=470)
        self.selfcomboExample.current(1)

        self.entry_res = Text(self, font=("Verdana", 10), width=39, height=8, wrap=WORD)
        self.entry_res.place(x=50, y=500)

        btn_load_img = tk.Button(self, text='Тест', command=self.test,font=("Verdana", 10),background = '#DCDCDC')
        btn_load_img.place(x=50, y=300)

        btn_test = tk.Button(self, text='Загрузить снимок', command=self.openfile,font=("Verdana", 10),
                             background = '#DCDCDC')
        btn_test.place(x=50, y=250)

        btn_cancel = tk.Button(self, text='Закрыть окно', command=self.destroy,font=("Verdana", 10),
                               background = '#DCDCDC')
        btn_cancel.place(x=262, y=640)

        btn_cancel = tk.Button(self, text='Сохранить диагноз', command=self.save_res, font=("Verdana", 10),
                               background='#DCDCDC')
        btn_cancel.place(x=50, y=640)

        self.grab_set()
        self.focus_set()

        self.ImageFrame = LabelFrame(self, height=555, width=620)
        self.ImageFrame.place(x=400,y=80)

    def save_res(self):
        dbCursor = connection.cursor()
        requestString = "UPDATE Test SET res = ?,usr_res= ?,res_r=? WHERE id= ? ;"
        dbCursor.execute(requestString, (self.entry_res.get("1.0",END),user.Surname+" "+user.Name,
                                         self.selfcomboExample.get(),self.id))
        connection.commit()
        app.view_records()
        bd.logr("Добавил диагноз")


    def openfile(self):
        self.filename = filedialog.askopenfilename(initialdir="D:\\Новая папка\\DataSets\\ALL\\IN",
                                                   title="Выбор снимка")
        self.ImageFrame.destroy()
        self.ImageFrame = LabelFrame(self, height=555, width=620)
        self.ImageFrame.place(x=400, y=80)
        self.our_image2 = Image.open(self.filename)
        self.our_image2 = resizeimage.resize_contain(self.our_image2, [555, 620])
        self.our_image2 = ImageTk.PhotoImage(self.our_image2)
        self.our_label2 = Label(self.ImageFrame,image=self.our_image2)
        self.our_label2.image = self.our_image2
        self.our_label2.pack()

    def test(self):
        self.regex = re.compile('(\D+)')
        self.regex2 = re.compile('\d{2}[-/]\d{2}[-/]\d{4}')
        self.regex3 = re.compile('[0-9]+')
        ree= self.regex3.findall("dddd")
        if self.entry_polis['fg'] == 'grey' or self.entry_data['fg']== 'grey' or \
            self.entry_name['fg']== 'grey' or self.entry_sorname['fg']== 'grey':
            self.label_err2.place_forget()
            self.label_err3.place_forget()
            self.label_err1.place(x=200, y=250)
        elif self.regex2.findall(self.entry_data.get())==ree:
             self.label_err1.place_forget()
             self.label_err3.place_forget()
             self.label_err2.place(x=200, y=250)
        elif self.regex.findall(self.entry_polis.get())!=ree:
             self.label_err1.place_forget()
             self.label_err2.place_forget()
             self.label_err3.place(x=200, y=250)
        else:
            self.label_err1.place_forget()
            self.label_err2.place_forget()
            self.label_err3.place_forget()
            img_path = self.filename
            img = image.load_img(img_path, target_size=(150, 150))
            x = image.img_to_array(img)
            x /= 255
            x = np.expand_dims(x, axis=0)
            predication = nei.final_model.predict(x)
            predication2 = nei.final_model.predict_classes(x)
            res = "Диагноз системы: "+self.classes[predication2[0]]+"\n\nУверенность системы: "+\
                  format(round(predication[0][predication2[0]]*100))+"%"
            res2 = self.classes[predication2[0]]
            res3 =format(round(predication[0][predication2[0]]*100))+"%"
            self.entry_sys_res.configure(state="normal")
            self.entry_sys_res.delete(1.0, END)
            self.entry_sys_res.insert(1.0, res)
            self.entry_sys_res.configure(state="disable")
            fin = open(self.filename, "rb")
            img = fin.read()
            fin.close()
            self.binary = Binary(img)
            dbCursor = connection.cursor()
            requestString = "insert into Patients(name,surname,patronymic,date_of_birth,polis) " \
                            "SELECT * FROM (SELECT ? as name,? as surname,? as patronomic,? as date_of_birth,? as polis) as tmp " \
                            "WHERE NOT EXISTS (SELECT polis FROM Patients WHERE polis = ?);"
            dbCursor.execute(requestString, (self.entry_sorname.get(),
                                             self.entry_name.get(),
                                             self.entry_patronymic.get(),
                                             self.entry_data.get(),
                                             self.entry_polis.get(),
                                             self.entry_polis.get()))

            requestString = "SELECT pat_id FROM Patients where polis=?"
            dbCursor.execute(requestString,self.entry_polis.get())
            for row in dbCursor:
                self.pat_id = row[0]
            requestString = "insert into Test(sys_res,date,usr_id,pat_id,sys_confidence) " \
                            "values(?,?,?,?,?);"
            dbCursor.execute(requestString, (res2,
                                             datetime.today().strftime('%Y-%m-%d'),
                                             user.id,
                                             self.pat_id,
                                             res3))
            requestString = "SELECT MAX(id) FROM Test"
            dbCursor.execute(requestString)
            for row in dbCursor:
                self.id = row[0]
            requestString = "insert into IMG_Data(img_dat,test_id) values(?,?);"
            dbCursor.execute(requestString, (self.binary, self.id))
            connection.commit()
            app.view_records()
            bd.logr("Произвёл диагностику")

class EntryWithPlaceholder(tk.Entry):
    def __init__(self, master=None, placeholder=None):
        super().__init__(master)

        if placeholder is not None:
            self.placeholder = placeholder
            self.placeholder_color = 'grey'
            self.default_fg_color = self['fg']

            self.bind("<FocusIn>", self.focus_in)
            self.bind("<FocusOut>", self.focus_out)

            self.put_placeholder()

    def put_placeholder(self):
        self.insert(0, self.placeholder)
        self['fg'] = self.placeholder_color

    def focus_in(self, *args):
        if self['fg'] == self.placeholder_color:
            self.delete('0', 'end')
            self['fg'] = self.default_fg_color

    def focus_out(self, *args):
        if not self.get():
            self.put_placeholder()

class neiron():

    def __init__(self):
        path = "D:\\Новая папка\\DataSets\\ALL\\IN"
        self.train_folder = path + '/train/'
        self.test_folder = path + '/test/'
        self.val_folder = path + '/val/'
        self.final_model = self.load()

    def learn(self):
        self.train_datagen = ImageDataGenerator(rescale=1. / 255,
                                           shear_range=0.2,
                                           zoom_range=0.2,
                                           horizontal_flip=True,
                                           rotation_range=40,
                                           width_shift_range=0.2,
                                           height_shift_range=0.2)

        self.test_datagen = ImageDataGenerator(rescale=1. / 255)

        self.training_set = self.train_datagen.flow_from_directory(self.train_folder,
                                                         target_size=(150, 150),
                                                         batch_size=32,
                                                         class_mode='categorical')

        self.val_set = self.test_datagen.flow_from_directory(self.val_folder,
                                                   target_size=(150, 150),
                                                   batch_size=32,
                                                   class_mode='categorical')

        self.test_set = self.test_datagen.flow_from_directory(self.test_folder,
                                                    target_size=(150, 150),
                                                    batch_size=32,
                                                    class_mode='categorical')

        self.model = Sequential()

        self.model.add(Conv2D(32, (3, 3), activation='relu', input_shape=(150, 150, 3), padding='same'))
        self.model.add(MaxPooling2D(pool_size=(2, 2)))

        self.model.add(Conv2D(32, (3, 3), activation='relu', padding='same'))
        self.model.add(MaxPooling2D(pool_size=(2, 2)))
        self.model.add(Dropout(0.25))

        self.model.add(Conv2D(64, (3, 3), activation='relu', padding='same'))
        self.model.add(MaxPooling2D(pool_size=(2, 2)))

        self.model.add(Conv2D(64, (3, 3), activation='relu', padding='same'))
        self.model.add(MaxPooling2D(pool_size=(2, 2)))
        self.model.add(Dropout(0.25))

        self.model.add(Conv2D(128, (3, 3), activation='relu', padding='same'))
        self.model.add(MaxPooling2D(pool_size=(2, 2)))

        self.model.add(Conv2D(128, (3, 3), activation='relu', padding='same'))
        self.model.add(MaxPooling2D(pool_size=(2, 2)))
        self.model.add(Dropout(0.25))

        self.model.add(Flatten())
        self.model.add(Dense(256, activation='relu'))
        self.model.add(Dropout(0.5))
        self.model.add(Dense(4, activation='sigmoid'))

        self.model.compile(optimizer='adam', loss='categorical_crossentropy', metrics=['accuracy'])

        model_train = self.model.fit_generator(self.training_set,
                                          steps_per_epoch=421,
                                          epochs=15,
                                          validation_data=self.val_set,
                                          validation_steps=6)

        test_accuracy = self.model.evaluate_generator(self.test_set, steps=31)
        print('Testing Accuracy: {:.2f}%'.format(test_accuracy[1] * 100))

    def save(self):
        self.model
        print("Сохраняем сеть")
        model_json = self.model.to_json()
        json_file = open("cifar10.json", "w")
        json_file.write(model_json)
        json_file.close()
        self.model.save_weights("cifar10.h5")
        print("Сохранение сети завершено")

    def load(self):
        json_file = open("cifar10.json", "r")
        loaded_model_json = json_file.read()
        json_file.close()
        loaded_model = model_from_json(loaded_model_json)
        loaded_model.load_weights("cifar10.h5")
        loaded_model.compile(loss='categorical_crossentropy', optimizer='adam', metrics=['accuracy'])
        return loaded_model

class User():
    def __init__(self,Name,Surname,id,admin):
        self.Name = Name
        self.Surname = Surname
        self.id = id
        self.admin = admin

class Bd():
    def logr(self,str):
        dbCursor = connection.cursor()
        requestString = "insert into Logr(action,act_date,usr_id) " \
                        "values(?,?,?);"
        dbCursor.execute(requestString, (str,
                                         datetime.today(),
                                         user.id))


class Authorization(tk.Frame):
    def __init__(self,aut):
        super().__init__(aut)
        self.auto = 0
        self.user = None
        self.init_main()

    def init_main(self):

        self.ImageFrame = LabelFrame(self, height=400, width=700)
        self.ImageFrame.pack()

        label_main = tk.Label(self, text='Авторизация:', font=("Verdana", 15, 'bold'))
        label_main.place(x=280, y=100)

        label_log = tk.Label(self, text='Логин:', font=("Verdana", 10))
        label_log.place(x=200, y=150)

        label_pass = tk.Label(self, text='Пароль:', font=("Verdana", 10))
        label_pass.place(x=200, y=180)

        self.label_err = tk.Label(self, text='Данные для входа ввдены не верно.\n Проверте правильность введёных данных',
                              font=("Verdana", 10), fg='#FF0000')

        self.entry_log = Entry(self,show="*")
        self.entry_log.configure(font=("Verdana", 10))
        self.entry_log.place(x=280, y=180)

        self.entry_pass = Entry(self)
        self.entry_pass.configure(font=("Verdana", 10))
        self.entry_pass.place(x=280, y=150)

        btn_cancel = tk.Button(self, text='Войти', command=self.check_log,font=("Verdana", 10),background = '#DCDCDC')
        btn_cancel.place(x=392, y=215)

    def check_log(self):
        dbCursor = connection.cursor()
        requestString = "select * from Usr where usr_login = ? and usr_password = ?"
        dbCursor.execute(requestString,(self.entry_pass.get(),self.entry_log.get()))
        row = dbCursor.fetchone()
        if row is not None:
            global user
            user = User(row[1],row[2],row[0],row[5])
            global b
            b = 1
            bd.logr("Авторизация")
            aut.destroy()
        else:
            self.label_err.place(x=210, y=250)


if __name__ == "__main__":
    bd = Bd()
    aut = tk.Tk()
    aut2 = Authorization(aut)
    aut2.pack()
    aut.title("Авторизация")
    aut.geometry("700x400+10+50")
    aut.resizable(False, False)
    global b
    b=0
    aut.mainloop()
    if(b == 1):
        root = tk.Tk()
        nei = neiron()
        app = Main(root)
        app.pack()
        root.title("Нейроэмулятор")
        root.geometry("1300x710+10+50")
        root.resizable(False, False)
        root.mainloop()